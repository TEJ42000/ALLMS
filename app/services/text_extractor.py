"""Unified Text Extraction Service.

Provides text extraction from various file types for LLM processing:
- PDFs (real PDFs using PyMuPDF)
- Slide Archives (ZIP with manifest.json containing JPEG slides + text)
- Images (PNG, JPG, etc. using OCR via pytesseract or Google Vision)
- Office Documents (DOCX using python-docx)
- Markdown/Text files (direct read)
- HTML files (using BeautifulSoup)

This service provides a consistent API for extracting text from any material
file, enabling future LLM functionality for quizzes, summaries, etc.
"""

import json
import logging
import mimetypes
from pathlib import Path
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, field

from app.services.syllabus_parser import validate_path_within_base, MATERIALS_BASE

# Alias for backward compatibility - MATERIALS_ROOT is used by text_cache_service
MATERIALS_ROOT = MATERIALS_BASE

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """Result of text extraction from a file."""
    file_path: str
    file_type: str
    text: str
    success: bool
    error: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    pages: Optional[List[Dict[str, Any]]] = None  # For multi-page documents


def get_file_extension(file_path: Path) -> str:
    """Get lowercase file extension without the dot."""
    return file_path.suffix.lower().lstrip('.')


def detect_file_type(file_path: Path) -> str:
    """Detect the actual file type, not just by extension.

    Returns one of: 'pdf', 'slide_archive', 'image', 'docx', 'markdown',
                    'text', 'html', 'json', 'unknown'
    """
    if not file_path.exists():
        return 'unknown'

    ext = get_file_extension(file_path)

    # Check for slide archive (ZIP with manifest.json disguised as PDF)
    if ext == 'pdf':
        from app.services.slide_archive import is_slide_archive, get_file_type as get_pdf_type
        pdf_type = get_pdf_type(file_path)
        if pdf_type == 'slide_archive':
            return 'slide_archive'
        elif pdf_type == 'pdf':
            return 'pdf'
        return 'unknown'

    # Image types
    if ext in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'):
        return 'image'

    # Office documents
    if ext == 'docx':
        return 'docx'
    if ext == 'pptx':
        return 'pptx'

    # Text-based formats
    if ext == 'md':
        return 'markdown'
    if ext == 'txt':
        return 'text'
    if ext == 'html' or ext == 'htm':
        return 'html'
    if ext == 'json':
        return 'json'

    return 'unknown'


def extract_text(file_path: Path | str, *, _skip_path_validation: bool = False) -> ExtractionResult:
    """Extract text from any supported file type.

    This is the main entry point for text extraction.

    Args:
        file_path: Path to the file (absolute or relative to project root)
        _skip_path_validation: Internal flag for testing only. DO NOT use in production code.
            When True, skips the path validation check. This is only for unit tests
            that need to test extraction from temporary directories.

    Returns:
        ExtractionResult with extracted text and metadata
    """
    # Security: Validate path to prevent path traversal attacks (CWE-22/23/36)
    # Path validation can be skipped for testing only (underscore prefix indicates internal use)
    if not _skip_path_validation:
        try:
            validated_path = validate_path_within_base(str(file_path), MATERIALS_BASE)
        except ValueError as e:
            logger.warning("Path validation failed for path=%s: %s", file_path, e)
            return ExtractionResult(
                file_path=str(file_path),
                file_type='unknown',
                text='',
                success=False,
                error=f'Invalid path: {file_path}'
            )
    else:
        validated_path = Path(file_path).resolve()

    if not validated_path.exists():
        return ExtractionResult(
            file_path=str(validated_path),
            file_type='unknown',
            text='',
            success=False,
            error=f'File not found: {validated_path}'
        )

    file_type = detect_file_type(validated_path)

    extractors = {
        'pdf': extract_from_pdf,
        'slide_archive': extract_from_slide_archive,
        'image': extract_from_image,
        'docx': extract_from_docx,
        'markdown': extract_from_text,
        'text': extract_from_text,
        'html': extract_from_html,
        'json': extract_from_json,
    }

    extractor = extractors.get(file_type)
    if not extractor:
        return ExtractionResult(
            file_path=str(validated_path),
            file_type=file_type,
            text='',
            success=False,
            error=f'Unsupported file type: {file_type}'
        )

    try:
        return extractor(validated_path)
    except Exception as e:
        logger.error("Text extraction failed for %s: %s", validated_path, e)
        return ExtractionResult(
            file_path=str(validated_path),
            file_type=file_type,
            text='',
            success=False,
            error=str(e)
        )


# ============================================================================
# Individual Extractors
# ============================================================================

def extract_from_pdf(file_path: Path) -> ExtractionResult:
    """Extract text from a real PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ExtractionResult(
            file_path=str(file_path),
            file_type='pdf',
            text='',
            success=False,
            error='PyMuPDF (fitz) not installed'
        )

    try:
        doc = fitz.open(str(file_path))
        pages = []
        all_text = []

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            pages.append({
                'page_number': page_num,
                'text': text,
                'char_count': len(text)
            })
            all_text.append(f"--- Page {page_num} ---\n{text}")

        doc.close()

        return ExtractionResult(
            file_path=str(file_path),
            file_type='pdf',
            text='\n'.join(all_text),
            success=True,
            metadata={'num_pages': len(pages)},
            pages=pages
        )
    except Exception as e:
        return ExtractionResult(
            file_path=str(file_path),
            file_type='pdf',
            text='',
            success=False,
            error=f'PDF extraction failed: {e}'
        )


def extract_from_slide_archive(file_path: Path) -> ExtractionResult:
    """Extract text from a slide archive (ZIP with JPEG slides + text files)."""
    from app.services.slide_archive import extract_slide_archive, get_all_text

    archive_data = extract_slide_archive(file_path)
    if not archive_data:
        return ExtractionResult(
            file_path=str(file_path),
            file_type='slide_archive',
            text='',
            success=False,
            error='Failed to extract slide archive'
        )

    # Build pages list
    pages = []
    for slide in archive_data.slides:
        pages.append({
            'page_number': slide.page_number,
            'text': slide.text_content or '',
            'char_count': len(slide.text_content or '')
        })

    all_text = get_all_text(file_path)

    return ExtractionResult(
        file_path=str(file_path),
        file_type='slide_archive',
        text=all_text,
        success=True,
        metadata={'num_pages': archive_data.num_pages},
        pages=pages
    )



def extract_from_image(file_path: Path) -> ExtractionResult:
    """Extract text from an image using OCR.

    Uses pytesseract if available, falls back to basic error message.
    For production, consider Google Cloud Vision API for better accuracy.
    """
    try:
        from PIL import Image
        import pytesseract
    except ImportError as e:
        return ExtractionResult(
            file_path=str(file_path),
            file_type='image',
            text='',
            success=False,
            error=f'OCR dependencies not installed: {e}. Install with: pip install pytesseract Pillow'
        )

    try:
        image = Image.open(file_path)

        # Run OCR
        text = pytesseract.image_to_string(image)

        # Get image metadata
        width, height = image.size

        return ExtractionResult(
            file_path=str(file_path),
            file_type='image',
            text=text.strip(),
            success=True,
            metadata={
                'width': width,
                'height': height,
                'format': image.format,
                'mode': image.mode
            }
        )
    except Exception as e:
        return ExtractionResult(
            file_path=str(file_path),
            file_type='image',
            text='',
            success=False,
            error=f'OCR extraction failed: {e}'
        )


def extract_from_docx(file_path: Path) -> ExtractionResult:
    """Extract text from a DOCX file using python-docx.

    Falls back to plain text extraction if the file is actually a text file
    with .docx extension.
    """
    import zipfile

    # First, check if it's actually a DOCX (ZIP archive) or just a text file
    try:
        with open(file_path, 'rb') as f:
            header = f.read(4)

        # DOCX files are ZIP archives and start with "PK" (0x504B)
        if header[:2] != b'PK':
            # Not a real DOCX, try reading as text
            logger.info("File %s appears to be text, not DOCX. Reading as text.", file_path)
            return extract_from_text(file_path)
    except Exception:
        pass  # Continue and let python-docx try

    try:
        from docx import Document
    except ImportError:
        return ExtractionResult(
            file_path=str(file_path),
            file_type='docx',
            text='',
            success=False,
            error='python-docx not installed. Install with: pip install python-docx'
        )

    try:
        doc = Document(str(file_path))

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return ExtractionResult(
            file_path=str(file_path),
            file_type='docx',
            text='\n\n'.join(paragraphs),
            success=True,
            metadata={
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
        )
    except Exception as e:
        # If python-docx fails, try reading as plain text
        logger.warning("DOCX extraction failed for %s, trying as text: %s", file_path, e)
        try:
            return extract_from_text(file_path)
        except Exception:
            return ExtractionResult(
                file_path=str(file_path),
                file_type='docx',
                text='',
                success=False,
                error=f'DOCX extraction failed: {e}'
            )


def extract_from_text(file_path: Path) -> ExtractionResult:
    """Extract text from plain text or markdown files."""
    try:
        text = file_path.read_text(encoding='utf-8')
        return ExtractionResult(
            file_path=str(file_path),
            file_type=get_file_extension(file_path),
            text=text,
            success=True,
            metadata={'char_count': len(text), 'line_count': text.count('\n') + 1}
        )
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            text = file_path.read_text(encoding='latin-1')
            return ExtractionResult(
                file_path=str(file_path),
                file_type=get_file_extension(file_path),
                text=text,
                success=True,
                metadata={'char_count': len(text), 'encoding': 'latin-1'}
            )
        except Exception as e:
            return ExtractionResult(
                file_path=str(file_path),
                file_type=get_file_extension(file_path),
                text='',
                success=False,
                error=f'Text file read failed: {e}'
            )
    except Exception as e:
        return ExtractionResult(
            file_path=str(file_path),
            file_type=get_file_extension(file_path),
            text='',
            success=False,
            error=f'Text file read failed: {e}'
        )


def extract_from_html(file_path: Path) -> ExtractionResult:
    """Extract text from HTML files."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        # Fallback: just strip tags with regex
        import re
        try:
            html = file_path.read_text(encoding='utf-8')
            text = re.sub('<[^<]+?>', '', html)
            return ExtractionResult(
                file_path=str(file_path),
                file_type='html',
                text=text.strip(),
                success=True,
                metadata={'method': 'regex_fallback'}
            )
        except Exception as e:
            return ExtractionResult(
                file_path=str(file_path),
                file_type='html',
                text='',
                success=False,
                error=f'HTML extraction failed: {e}'
            )

    try:
        html = file_path.read_text(encoding='utf-8')
        soup = BeautifulSoup(html, 'html.parser')

        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()

        text = soup.get_text(separator='\n', strip=True)

        return ExtractionResult(
            file_path=str(file_path),
            file_type='html',
            text=text,
            success=True,
            metadata={'title': soup.title.string if soup.title else None}
        )
    except Exception as e:
        return ExtractionResult(
            file_path=str(file_path),
            file_type='html',
            text='',
            success=False,
            error=f'HTML extraction failed: {e}'
        )


def extract_from_json(file_path: Path) -> ExtractionResult:
    """Extract text content from JSON files."""
    try:
        text = file_path.read_text(encoding='utf-8')
        data = json.loads(text)

        # For JSON, we return both the raw JSON and attempt to extract text values
        def extract_strings(obj, depth=0):
            """Recursively extract string values from JSON."""
            if depth > 10:  # Prevent infinite recursion
                return []
            strings = []
            if isinstance(obj, str):
                if len(obj) > 10:  # Only meaningful strings
                    strings.append(obj)
            elif isinstance(obj, dict):
                for v in obj.values():
                    strings.extend(extract_strings(v, depth + 1))
            elif isinstance(obj, list):
                for item in obj:
                    strings.extend(extract_strings(item, depth + 1))
            return strings

        extracted_strings = extract_strings(data)

        return ExtractionResult(
            file_path=str(file_path),
            file_type='json',
            text='\n'.join(extracted_strings) if extracted_strings else text,
            success=True,
            metadata={
                'is_structured': True,
                'string_count': len(extracted_strings)
            }
        )
    except Exception as e:
        return ExtractionResult(
            file_path=str(file_path),
            file_type='json',
            text='',
            success=False,
            error=f'JSON extraction failed: {e}'
        )


# ============================================================================
# Batch Operations
# ============================================================================

def extract_all_from_folder(
    folder_path: Path,
    recursive: bool = True,
    *,
    _skip_path_validation: bool = False
) -> List[ExtractionResult]:
    """Extract text from all supported files in a folder.

    Args:
        folder_path: Path to the folder
        recursive: Whether to search subdirectories
        _skip_path_validation: Internal flag for testing only. DO NOT use in production code.

    Returns:
        List of ExtractionResult for each file
    """
    results = []

    if not folder_path.exists():
        return results

    pattern = '**/*' if recursive else '*'

    for file_path in folder_path.glob(pattern):
        if file_path.is_file():
            # Check if it's a supported type
            file_type = detect_file_type(file_path)
            if file_type \!= 'unknown':
                result = extract_text(file_path, _skip_path_validation=_skip_path_validation)
                results.append(result)

    return results


def get_extraction_summary(results: List[ExtractionResult]) -> Dict[str, Any]:
    """Get a summary of extraction results.

    Args:
        results: List of ExtractionResult objects

    Returns:
        Summary dict with counts and statistics
    """
    successful = [r for r in results if r.success]
    failed = [r for r in results if not r.success]

    by_type = {}
    for r in results:
        if r.file_type not in by_type:
            by_type[r.file_type] = {'success': 0, 'failed': 0, 'total_chars': 0}
        if r.success:
            by_type[r.file_type]['success'] += 1
            by_type[r.file_type]['total_chars'] += len(r.text)
        else:
            by_type[r.file_type]['failed'] += 1

    return {
        'total_files': len(results),
        'successful': len(successful),
        'failed': len(failed),
        'total_characters': sum(len(r.text) for r in successful),
        'by_type': by_type,
        'errors': [{'file': r.file_path, 'error': r.error} for r in failed]
    }
